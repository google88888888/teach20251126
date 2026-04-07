import pandas as pd
from datetime import datetime
import json
from openpyxl import load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows
import xlwings as xw
import shutil
import requests
import os
import openpyxl
from docx import Document
from collections import deque
from docx.shared import Pt
from docx.oxml.ns import qn
from pyecharts.charts import Bar
from pyecharts import options as opts
import re
import os
import win32com.client

import os
from docx import Document

def get_first_tables(folder_path):
    tables_list = []

    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            
            try:
                doc = Document(file_path)
                
                if doc.tables:  # 有表格
                    first_table = doc.tables[0]
                    
                    # 转成二维列表（更好用）
                    table_data = []
                    for row in first_table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    
                    tables_list.append(table_data)

            except Exception as e:
                print(f"处理失败: {filename}, 错误: {e}")

    return tables_list


# 使用示例
folder = f"./customOrder"
all_tables = get_first_tables(folder)


with open('all_tables.json', 'w', encoding='utf-8') as f:
    json.dump(all_tables, f, ensure_ascii=False, indent=4)

all_tables_with_simple_origin_data=[]
for index,item in enumerate(all_tables):
    all_tables_with_simple_origin_data.append({
        '客户名称':re.sub(r"\s+", "", item[2][1]),
        '存款金额':re.sub(r"\s+", "", item[3][1]),
        '起息日':re.sub(r"\s+", "", item[3][3]),
        '到期日':re.sub(r"\s+", "", item[4][1]),
        '对客高收益报价':re.sub(r"\s+", "", item[8][1]),
    })
with open('all_tables_with_simple_origin_data.json', 'w', encoding='utf-8') as f:
    json.dump(all_tables_with_simple_origin_data, f, ensure_ascii=False, indent=4)



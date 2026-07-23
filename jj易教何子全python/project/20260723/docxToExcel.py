import os
import re
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Border, Side, numbers, Font, Alignment
from docx.oxml.ns import qn
import json

def try_convert_to_number(s):
    """将可能含千分位逗号的字符串转为数字，成功返回(数字,True)，否则返回(None,False)"""
    if not isinstance(s, str):
        return s, False
    s = s.strip()
    if s == "":
        return None, False
    s_clean = s.replace(',', '')
    if re.match(r'^-?\d+(?:\.\d+)?$', s_clean):
        try:
            num = float(s_clean)
            if num.is_integer():
                return int(num), True
            else:
                return num, True
        except ValueError:
            return None, False
    return None, False

def paragraph_has_bold(para):
    """判断段落中是否有加粗的文本"""
    for run in para.runs:
        if run.bold:
            return True
    return False

def get_word_cell_vertical_alignment(cell):
    """读取Word单元格的垂直对齐方式，返回 'top'/'center'/'bottom' 或 None"""
    tcPr = cell._element.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    v_align = tcPr.find(qn('w:vAlign'))
    if v_align is None:
        return None
    val = v_align.get(qn('w:val'))
    if val in ('top', 'center', 'bottom'):
        return val
    return None

def get_word_para_horizontal_alignment(para):
    """读取Word段落的水平对齐方式，返回 'left'/'center'/'right'/'justify' 或 None"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        return None
    val = jc.get(qn('w:val'))
    mapping = {
        'left': 'left', 'start': 'left',
        'center': 'center', 'centered': 'center',
        'right': 'right', 'end': 'right',
        'both': 'justify', 'justify': 'justify',
        'distribute': 'justify',
    }
    return mapping.get(val)

def get_word_cell_horizontal_alignment(cell):
    """读取Word单元格的水平对齐方式（取第一个有对齐设置的段落）"""
    for para in cell.paragraphs:
        h_align = get_word_para_horizontal_alignment(para)
        if h_align is not None:
            return h_align
    return None

def get_cell_merge_info(table):
    """获取Word表格中的合并单元格信息，返回应该被合并的单元格范围列表"""
    merge_info = []
    processed_cells = set()
    
    rows = table.rows
    cols = len(rows[0].cells) if rows else 0
    
    for r, row in enumerate(rows):
        for c, cell in enumerate(row.cells):
            if (r, c) in processed_cells:
                continue
            
            # 检查单元格的tcPr属性，查看是否有合并信息
            tcPr = cell._element.get_or_add_tcPr()
            
            # 检查水平合并（gridSpan）
            grid_span = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
            h_span = int(grid_span.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')) if grid_span is not None else 1
            
            # 检查垂直合并（vMerge）
            v_merge = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')     
            # 对于垂直合并，需要找到合并的范围
            v_span = 1
            if v_merge is not None:
                # 检查是否是vMerge的restart（合并的开始），如果是合并的开始，计算垂直跨度
                for check_r in range(r + 1, len(rows)):
                    check_cell = rows[check_r].cells[c] if c < len(rows[check_r].cells) else None
                    if check_cell is None:
                        break
                    check_tcPr = check_cell._element.get_or_add_tcPr()
                    check_v_merge = check_tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
                    if check_v_merge is not None :   
                        v_span += 1
                    else:
                        break
            
            if h_span > 1 or v_span > 1:
                merge_info.append((r, c, v_span, h_span))
                for mr in range(r, r + v_span):
                    for mc in range(c, c + h_span):
                        processed_cells.add((mr, mc))
    return merge_info

def extract_with_formatting(docx_path, target_title, output_excel):
    doc = Document(docx_path)
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    found = False
    row_idx = 1
    table_start_rows = []  # 记录每个表格在Excel中的起始行、合并信息、最大列数和结束行

    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    for element in doc.element.body:
        if element.tag.endswith('p'):          # 段落
            para = next((p for p in doc.paragraphs if p._element is element), None)
            if para is None:
                continue
            text = para.text.strip()
            if not found:
                if target_title in text:
                    found = True
                else:
                    continue
            cell = ws.cell(row=row_idx, column=1, value=text)
            # 如果段落中有加粗文本，则设置单元格字体加粗
            if paragraph_has_bold(para):
                cell.font = Font(bold=True)
            row_idx += 1

        elif element.tag.endswith('tbl'):       # 表格
            if not found:
                continue
            table = next((t for t in doc.tables if t._element is element), None)
            if table is None:
                continue

            start_row = row_idx
            merge_info = get_cell_merge_info(table)
            
            # 写入表格数据，同时处理数字格式和加粗
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    raw_text = cell.text.strip()
                    num_val, is_num = try_convert_to_number(raw_text)
                    excel_cell = ws.cell(row=start_row + r, column=1 + c)
                    if is_num:
                        excel_cell.value = num_val
                        if isinstance(num_val, int):
                            excel_cell.number_format = '#,##0'
                        else:
                            excel_cell.number_format = '#,##0.00'
                    else:
                        excel_cell.value = raw_text
                    
                    # 检查单元格中是否有加粗文本，如果有则设置 Excel 单元格为加粗
                    has_bold = False
                    for para in cell.paragraphs:
                        if paragraph_has_bold(para):
                            has_bold = True
                            break
                    if has_bold:
                        excel_cell.font = Font(bold=True)
                    # 保持单元格的水平和垂直对齐方式
                    h_align = get_word_cell_horizontal_alignment(cell)
                    v_align = get_word_cell_vertical_alignment(cell)
                    if h_align is not None or v_align is not None:
                        excel_cell.alignment = Alignment(horizontal=h_align, vertical=v_align)
            # 加边框：顶部/底部加粗，左右边框仅内部细线，最左列左边框无，最右列右边框无
            end_row = start_row + len(table.rows) - 1
            max_col = max((len(row.cells) for row in table.rows), default=0)
            table_start_rows.append((start_row, end_row, max_col, merge_info))
            
            # 创建一个字典来快速查找合并信息
            merge_dict = {(r, c): (v_span, h_span) for r, c, v_span, h_span in merge_info}
            
            top_bold = Side(style='medium')
            bottom_bold = Side(style='medium')
            thin_side = Side(style='thin')
            no_side = Side(style=None)
            
            for i in range(start_row, end_row + 1):
                for j in range(1, max_col + 1):
                    # 获取相对表格的行列位置
                    rel_row = i - start_row
                    rel_col = j - 1
                    
                    # 检查这个单元格是否是合并区域的一部分（但不是起始单元格）
                    is_merged_non_start = False
                    for r, c, v_span, h_span in merge_info:
                        if r <= rel_row < r + v_span and c <= rel_col < c + h_span and not (rel_row == r and rel_col == c):
                            is_merged_non_start = True
                            break
                    
                    if is_merged_non_start:
                        continue  # 跳过合并区域内的非起始单元格
                    
                    # 检查这个单元格是否是合并起始单元格
                    if (rel_row, rel_col) in merge_dict:
                        v_span, h_span = merge_dict[(rel_row, rel_col)]
                        # 对于合并单元格，根据合并范围设置边框
                        top = top_bold if i == start_row else thin_side
                        bottom = bottom_bold if i + v_span - 1 == end_row else thin_side
                        left = no_side if j == 1 else thin_side
                        right = no_side if j + h_span - 1 == max_col else thin_side
                    else:
                        # 普通单元格的边框
                        top = top_bold if i == start_row else thin_side
                        bottom = bottom_bold if i == end_row else thin_side
                        left = no_side if j == 1 else thin_side
                        right = no_side if j == max_col else thin_side
                    
                    ws.cell(row=i, column=j).border = Border(
                        left=left,
                        right=right,
                        top=top,
                        bottom=bottom
                    )
            # 表格后空一行
            row_idx += len(table.rows) + 1

    # 处理合并单元格
    for start_row, end_row, max_col, merge_info in table_start_rows:
        for r, c, v_span, h_span in merge_info:
            if v_span > 1 or h_span > 1:
                start_cell = ws.cell(row=start_row + r, column=1 + c)
                end_cell = ws.cell(row=start_row + r + v_span - 1, column=1 + c + h_span - 1)
                ws.merge_cells(f'{start_cell.coordinate}:{end_cell.coordinate}')

    wb.save(output_excel)
    print(f"已保存至: {output_excel}")

if __name__ == "__main__":
    docx_file = r"./3、佑荣科技2025年财审报告附注.docx"
    excel_file = r"./附注_合并项目注释_完整格式1.xlsx"
    # os.makedirs(os.path.dirname(excel_file), exist_ok=True)
    extract_with_formatting(docx_file, "八、财务报表主要项目注释", excel_file)



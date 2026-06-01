from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook


def is_blank_row(row_values):
    return all(value is None or str(value).strip() == "" for value in row_values)


def format_value_for_doc(value):
    if value is None:
        return ""
    if isinstance(value, bool):
        return str(value)
    if isinstance(value, (int, float)):
        return f"{value:,.2f}"
    text = str(value).strip()
    if text == "":
        return ""
    try:
        num = float(text.replace(",", ""))
        return f"{num:,.2f}"
    except ValueError:
        return text


def set_paragraph_font(para):
    for run in para.runs:
        run.font.name = "宋体"
        run.font.size = Pt(10.5)
        rfonts = run._element.rPr.rFonts
        rfonts.set(qn("w:eastAsia"), "宋体")


def set_cell_font(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = "宋体"
            run.font.size = Pt(10.5)
            rfonts = run._element.rPr.rFonts
            rfonts.set(qn("w:eastAsia"), "宋体")


def iter_doc_elements_from_title(doc, target_title):
    started = False
    for child in doc.element.body:
        if child.tag.endswith("p"):
            para = next((p for p in doc.paragraphs if p._element is child), None)
            if para is None:
                continue
            if not started and target_title in para.text:
                started = True
                yield "p", para
            elif started:
                yield "p", para
        elif started and child.tag.endswith("tbl"):
            table = next((t for t in doc.tables if t._element is child), None)
            if table is not None:
                yield "tbl", table


def load_excel_rows(excel_path):
    wb = load_workbook(excel_path, data_only=True)
    ws = wb.active
    # 构建包含工作表所有单元格值和加粗标志的矩阵
    max_row = ws.max_row
    max_col = ws.max_column
    grid_values = [[ws.cell(row=r, column=c).value for c in range(1, max_col + 1)] for r in range(1, max_row + 1)]
    grid_bolds = [[bool(ws.cell(row=r, column=c).font and ws.cell(row=r, column=c).font.bold) for c in range(1, max_col + 1)] for r in range(1, max_row + 1)]

    # 展开所有合并单元格区域：将合并区域内每个格子都填充为主单元格（左上角）的值和加粗标志
    for mrange in ws.merged_cells.ranges:
        try:
            min_row, min_col, max_row_m, max_col_m = mrange.min_row, mrange.min_col, mrange.max_row, mrange.max_col
        except Exception:
            # 兼容老版本 openpyxl 的字符串表示
            bounds = mrange.bounds  # (min_col, min_row, max_col, max_row) in some versions
            min_col, min_row, max_col_m, max_row_m = bounds
        val = ws.cell(row=min_row, column=min_col).value
        bold_flag = bool(ws.cell(row=min_row, column=min_col).font and ws.cell(row=min_row, column=min_col).font.bold)
        for rr in range(min_row, max_row_m + 1):
            for cc in range(min_col, max_col_m + 1):
                grid_values[rr - 1][cc - 1] = val
                grid_bolds[rr - 1][cc - 1] = bold_flag

    # 返回值和加粗信息的并行矩阵（行作为 tuple）
    return [tuple(row) for row in grid_values], [tuple(row) for row in grid_bolds]


def update_doc_from_excel(doc_path, excel_path, output_path, target_title):
    doc = Document(doc_path)
    excel_rows, excel_bolds = load_excel_rows(excel_path)
    excel_index = 0

    def next_nonempty_row():
        nonlocal excel_index
        while excel_index < len(excel_rows) and is_blank_row(excel_rows[excel_index]):
            excel_index += 1
        if excel_index >= len(excel_rows):
            return None
        row = excel_rows[excel_index]
        bolds = excel_bolds[excel_index]
        excel_index += 1
        return row, bolds

    for elem_type, elem in iter_doc_elements_from_title(doc, target_title):
        if elem_type == "p":
            nr = next_nonempty_row()
            if nr is None:
                break
            row_vals, row_bolds = nr
            value = row_vals[0] if len(row_vals) > 0 else None
            bold_flag = row_bolds[0] if len(row_bolds) > 0 else False
            text = format_value_for_doc(value)
            # 清除原有内容并以新的 run 写入，保持加粗信息
            elem.text = ""
            run = elem.add_run(text)
            run.bold = bool(bold_flag)
            run.font.name = "宋体"
            run.font.size = Pt(10.5)
            try:
                rfonts = run._element.rPr.rFonts
                rfonts.set(qn("w:eastAsia"), "宋体")
            except Exception:
                pass
        elif elem_type == "tbl":
            first_table_row = True
            for table_row in elem.rows:
                nr = next_nonempty_row()
                if nr is None:
                    break
                row_vals, row_bolds = nr
                for col_idx, cell in enumerate(table_row.cells):
                    value = row_vals[col_idx] if col_idx < len(row_vals) else None
                    bold_flag = row_bolds[col_idx] if col_idx < len(row_bolds) else False
                    text = format_value_for_doc(value)
                    # debug: 输出 Excel 中的值与加粗标志，便于排查
                    if first_table_row:
                        print(f"[DEBUG] table header col={col_idx} excel_value={repr(value)} bold={bold_flag}")
                    # 清空 cell 并写入带 run 的段落以设置加粗
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(text)
                    run.bold = bool(bold_flag)
                    run.font.name = "宋体"
                    run.font.size = Pt(10.5)
                    try:
                        rfonts = run._element.rPr.rFonts
                        rfonts.set(qn("w:eastAsia"), "宋体")
                    except Exception:
                        pass
                first_table_row = False
            while excel_index < len(excel_rows) and is_blank_row(excel_rows[excel_index]):
                excel_index += 1

    doc.save(output_path)
    print(f"已保存回 Word：{output_path}")


if __name__ == "__main__":
    # 直接使用固定路径（如需修改请在此处更改）
    docx_path = Path("./3、佑荣科技2025年财审报告附注.docx")
    excel_path = Path("./附注_合并项目注释_完整格式1.xlsx")
    output_path = Path("./3、佑荣科技2025年财审报告附注_更新.docx")
    title = "八、财务报表主要项目注释"

    update_doc_from_excel(docx_path, excel_path, output_path, title)

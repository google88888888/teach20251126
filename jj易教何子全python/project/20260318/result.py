# 使用天眼查的数据
# https://open.tianyancha.com/open/451
import pandas as pd
from datetime import datetime
import json
from openpyxl import load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows
import xlwings as xw
import shutil
import requests
import os
import openpyxl
from docx import Document
from collections import deque
from docx.shared import Pt
from docx.oxml.ns import qn

word_path = "1.docx"
doc = Document(word_path)
tables = doc.tables

CACHE_FILE = "cache.json"

# 启动时加载缓存
def load_cache():
    if os.path.exists(CACHE_FILE):
        with open(CACHE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

# 保存缓存到文件
def save_cache(cache):
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)


# 初始化缓存（全局）
cache = load_cache()

def fetch_data(url):

    print('url',url)

    if url in cache:
        print("走缓存")
        return cache[url]

    print("调用接口")
    headers={'Authorization': "28041b0a-feec-49df-9d42-21c21e208bca"}
    res = requests.get(url, headers=headers)
    response=res.json()

    cache[url] = response

    # 每次更新都写入文件（简单但安全）
    save_cache(cache)

    return response

res = fetch_data("http://open.api.tianyancha.com/services/v3/open/investtree?flag=4&dir=down&keyword=深圳市前海一方科技研发集团有限公司&minPercent=0&maxPercent=1")
# 有些接口的result是JSON字符串,有些接口的result是JSON对象,要区别处理
investtree = json.loads(res["result"])

rankToChinese={
    1:'一级',
    2:'二级',
    3:'三级',
    4:'四级',
    5:'五级',
    6:'六级',
    7:'七级',
    8:'八级',
    9:'九级',
    10:'十级',
}

with open('investtree.json', 'w', encoding='utf-8') as f:
    json.dump(investtree, f, ensure_ascii=False, indent=4)

queue = deque([(investtree[0], 0)])
investtreeList = []
while queue:
    node, level = queue.popleft()
    if(level!=0 and node['regStatus']=="存续"):
        currentType='控股公司'
        if level == 1:
            currentType='控股子公司'
        percent=f"{float(node['percent'])*100:.2f}"
        investtreeList.append({
            '子公司名称':node['name'],
            '子公司类型':currentType,
            '级次':rankToChinese[level],
            '持股比例（%）':percent,
            '表决权比例（%）':percent,
            '状态':node['regStatus'],
            '社会统一信用代码':node['creditCode'],
        })

    for index,item in enumerate(node['children']):
        queue.append((item, level + 1))

for index,item in enumerate(investtreeList):
    currentRes=fetch_data(f"http://open.api.tianyancha.com/services/open/ic/baseinfoV3/2.0?keyword={item['社会统一信用代码']}")
    item['注册地']=currentRes['result']['regLocation']
    item['业务性质']=currentRes['result']['industryAll']['category']
    item['对合营企业或联营企业投资的会计处理方法']='权益法'
    item['注册资本（万元）']=currentRes['result']['regCapital'].split("万")[0]
    timestamp_ms = currentRes['result']['estiblishTime']
    timestamp_s = timestamp_ms / 1000
    dt = datetime.fromtimestamp(timestamp_s)
    date_str = dt.strftime('%Y-%m-%d')
    item['成立日期（注册日期）']=date_str

with open('investtreeList.json', 'w', encoding='utf-8') as f:
    json.dump(investtreeList, f, ensure_ascii=False, indent=4)

for i in range(len(tables[1].rows) - 1, 0, -1):
    tbl = tables[1]._tbl
    tr = tables[1].rows[i]._tr
    tbl.remove(tr)

for i in range(len(tables[2].rows) - 1, 1, -1):
    tbl = tables[2]._tbl
    tr = tables[2].rows[i]._tr
    tbl.remove(tr)

for i in range(len(tables[3].rows) - 1, 0, -1):
    tbl = tables[3]._tbl
    tr = tables[3].rows[i]._tr
    tbl.remove(tr)

for index,item in enumerate(investtreeList):
    dateLimitString = "2025-12-31"
    dateLimit = datetime.strptime(dateLimitString, "%Y-%m-%d")

    dateItemString = item['成立日期（注册日期）']
    dateItem = datetime.strptime(dateItemString, "%Y-%m-%d")

    if dateItem>dateLimit:
        continue
    needHumanLook=False
    changeinfoRes=fetch_data(f"http://open.api.tianyancha.com/services/open/ic/changeinfo/2.0?keyword={item['子公司名称']}&pageNum=1&pageSize=20")
    for changeinfoResItem in changeinfoRes['result']['items']:
        changeTimeString = changeinfoResItem['changeTime']
        changeTime = datetime.strptime(changeTimeString, "%Y-%m-%d")
        if changeinfoResItem['changeItem']=='投资人变更（包括出资额、出资方式、出资日期、投资人名称等）' and changeTime>dateLimit:
            needHumanLook=True
            break
    if needHumanLook==True:
        item['变更记录']=changeinfoRes['result']['items']
        with open(f'{item['子公司名称']}的信息及变更记录.json', 'w', encoding='utf-8') as f:
            json.dump(investtreeList, f, ensure_ascii=False, indent=4)
        print(f'{item['子公司名称']}的"投资人变更（包括出资额、出资方式、出资日期、投资人名称等）"在{dateLimitString}之后有变动，请人工查看{item['子公司名称']}的信息及变更记录.json')
        continue

    if float(item['持股比例（%）'])<40:
        row = tables[2].add_row().cells
        valueList=[
            item['子公司名称'],
            item['注册地'],
            item['业务性质'],
            item['持股比例（%）'],
            '',
            item['表决权比例（%）'],
            item['对合营企业或联营企业投资的会计处理方法'],
            item['注册资本（万元）']
        ]
        for indexValueList,itemValueList in enumerate(valueList):
            p = row[indexValueList].paragraphs[0]
            run = p.add_run(itemValueList)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    else:
        row = tables[1].add_row().cells
        valueList=[
            item['子公司名称'],
            item['子公司类型'],
            item['级次'],
            item['持股比例（%）'],
            item['表决权比例（%）'],
        ]
        for indexValueList,itemValueList in enumerate(valueList):
            p = row[indexValueList].paragraphs[0]
            run = p.add_run(itemValueList)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

doc.save("1_填充后.docx")

realControlPerson={
    'name':'深圳市前海一方科技研发集团有限公司',
    'humanName':'赖少丽',
}
# 我的思路是：从深圳市前海一方科技研发集团有限公司的法定代表人赖少丽入手，查他的“人员所有合作伙伴”，然后对于所有的合作伙伴分别查每一个人的“人员控股企业”
currentRes=fetch_data(f"http://open.api.tianyancha.com/services/v4/open/partners?name={realControlPerson['name']}&humanName={realControlPerson['humanName']}")
allPartner = {}
for item in currentRes['result']['items']:
    if item['hid'] in allPartner:
        allPartner[item['hid']] = {
            'cid':item['cid'],
            'hid':item['hid'],
            'name':item['name'],
            'count':allPartner[item['hid']]['count']+1
        }
    else:
        allPartner[item['hid']] = {
            'cid':item['cid'],
            'hid':item['hid'],
            'name':item['name'],
            'count':1
        }
    for itemPartner in item['partners']:
        if itemPartner['hid'] in allPartner:
            allPartner[itemPartner['hid']] = {
                'cid':itemPartner['cid'],
                'hid':itemPartner['hid'],
                'name':itemPartner['name'],
                'count':allPartner[itemPartner['hid']]['count']+1
            }
        else:
            allPartner[itemPartner['hid']] = {
                'cid':itemPartner['cid'],
                'hid':itemPartner['hid'],
                'name':itemPartner['name'],
                'count':1
            }
with open('allPartner.json', 'w', encoding='utf-8') as f:
    json.dump(allPartner, f, ensure_ascii=False, indent=4)

allPartnerOrder = list(allPartner.values())
allPartnerOrder.sort(key=lambda x: x["count"], reverse=True)
with open('allPartnerOrder.json', 'w', encoding='utf-8') as f:
    json.dump(allPartnerOrder, f, ensure_ascii=False, indent=4)

realControlPersonCompany=[]
realControlPersonCompanyOnlyName=[]

for item in allPartnerOrder:
    if item['name']==realControlPerson['humanName']:
        pageNum=1
        currentPartnerCompany=[]
        while True:
            currentRes=fetch_data(f"http://open.api.tianyancha.com/services/open/human/companyholding/2.0?hid={item['hid']}&pageSize=20&pageNum={pageNum}&cid={item['cid']}")
            pageNum=pageNum+1
            for itemCompany in ((currentRes or {}).get('result') or {}).get('items') or []:
                currentPartnerCompany.append(itemCompany['name'])
                if itemCompany['name'] not in realControlPersonCompanyOnlyName:
                    realControlPersonCompany.append([
                        itemCompany['name'],
                        '实际控制人控制的公司'
                    ])
                    realControlPersonCompanyOnlyName.append(itemCompany['name'])
            if len(currentPartnerCompany)>=(((currentRes or {}).get('result') or {}).get('total') or 0):
                break
        break

for item in allPartnerOrder:
    if item['name']!=realControlPerson['humanName'] and item['count']>=5:
        pageNum=1
        currentPartnerCompany=[]
        while True:
            currentRes=fetch_data(f"http://open.api.tianyancha.com/services/open/human/companyholding/2.0?hid={item['hid']}&pageSize=20&pageNum={pageNum}&cid={item['cid']}")
            pageNum=pageNum+1
            for itemCompany in ((currentRes or {}).get('result') or {}).get('items') or []:
                currentPartnerCompany.append(itemCompany['name'])
                if itemCompany['name'] not in realControlPersonCompanyOnlyName:
                    realControlPersonCompany.append([
                        itemCompany['name'],
                        '实际控制人关系亲密的家庭成员控制的公司'
                    ])
                    realControlPersonCompanyOnlyName.append(itemCompany['name'])
            if len(currentPartnerCompany)>=(((currentRes or {}).get('result') or {}).get('total') or 0):
                break


for index,item in enumerate(realControlPersonCompany):
    row = tables[3].add_row().cells
    for indexValueList,itemValueList in enumerate(item):
        p = row[indexValueList].paragraphs[0]
        run = p.add_run(itemValueList)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

doc.save("1_填充后.docx")